import os
import tempfile
from typing import Optional
import PyPDF2
from docx import Document
from pathlib import Path

class FileProcessor:
    """Service for processing and extracting text from uploaded files"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx'}
    
    async def extract_text(self, file_path: str, filename: str) -> str:
        """
        Extract text content from uploaded file
        
        Args:
            file_path: Path to the temporary file
            filename: Original filename to determine file type
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            Exception: If file processing fails
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                return await self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise Exception("PDF is password-protected. Please provide an unprotected version.")
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        text_content.append(page_text)
                
                if not text_content:
                    raise Exception("No text could be extracted from the PDF. The file may be image-based or corrupted.")
                
                return '\n\n'.join(text_content)
                
        except PyPDF2.errors.PdfReadError as e:
            raise Exception(f"Invalid PDF file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise Exception("No text could be extracted from the DOCX file.")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 10) -> bool:
        """
        Validate file size is within limits
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in MB
            
        Returns:
            bool: True if file size is valid
        """
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            return file_size <= max_size_bytes
        except OSError:
            return False
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get information about the uploaded file
        
        Args:
            file_path: Path to the file
            
        Returns:
            dict: File information including size, type, etc.
        """
        try:
            stat = os.stat(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            return {
                'size_bytes': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'extension': file_extension,
                'is_supported': file_extension in self.supported_extensions,
                'modified_time': stat.st_mtime
            }
        except OSError as e:
            return {
                'error': f"Could not get file info: {str(e)}"
            }
